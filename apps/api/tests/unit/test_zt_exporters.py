"""ZT PDF + XLSX exporter smokes for both frameworks."""

from __future__ import annotations

import io
import uuid

import pytest

from app.models.zt_assessment import (
    ZtAnswer,
    ZtAssessment,
    ZtAssessmentStatus,
    ZtFramework,
)
from app.zt.catalog import capabilities
from app.zt.exporters import build_context, render_docx, render_pdf, render_xlsx
from app.zt.maturity import ZtFrameworkCode, level_count
from app.zt.scoring import analyze_gaps, compute


def _build_inputs(
    framework: ZtFrameworkCode, *, stage: int | None = 3
) -> tuple[ZtAssessment, list[ZtAnswer]]:
    db_framework = (
        ZtFramework.CISA_ZTMM_2_0
        if framework == ZtFrameworkCode.CISA_ZTMM_2_0
        else ZtFramework.DOD_ZTRA
    )
    a = ZtAssessment(
        id=uuid.uuid4(),
        service_id=uuid.uuid4(),
        framework=db_framework,
        version=1,
        status=ZtAssessmentStatus.APPROVED,
    )
    answers: list[ZtAnswer] = []
    for cap in capabilities(framework):
        answers.append(
            ZtAnswer(
                id=uuid.uuid4(),
                assessment_id=a.id,
                capability_code=cap.code,
                maturity_stage=stage,
            )
        )
    return a, answers


def _ctx(framework: ZtFrameworkCode, stage: int | None = 3, *, target: int | None = None):
    """Build an exporter context. `target` defaults to the FRAMEWORK'S ceiling.

    It used to default to a hardcoded 4, and DoD ZTRA has three stages -- so
    `test_dod_xlsx_renders`, `test_dod_xlsx_answers_sheet_row_count` and
    `test_dod_pdf_renders` had been rendering DoD deliverables against a target
    the framework does not have, for as long as they had existed. Nothing
    failed, because `analyze_gaps` silently clamped 4 to 3 (#125). The clamp is
    now a refusal, which is what surfaced these three.

    Deriving the default from `level_count` rather than restating a number per
    framework means the fixture CANNOT express an impossible target by
    accident again. Callers that want a specific target still pass one, and
    passing an out-of-range one still raises -- that is the engine contract,
    pinned in `test_zt_target_stage_provenance.py`.
    """
    if target is None:
        target = level_count(framework)
    a, answers = _build_inputs(framework, stage=stage)
    stage_map = {ans.capability_code: ans.maturity_stage for ans in answers}
    score = compute(framework, stage_map)
    gap = analyze_gaps(framework, stage_map, target_stage=target)
    return build_context(
        client_legal_name="Atlas Defense Solutions",
        service_title="Zero Trust Assessment",
        framework=framework,
        assessment=a,
        answers=answers,
        score=score,
        gap=gap,
    )


@pytest.mark.unit
def test_cisa_xlsx_has_three_sheets() -> None:
    from openpyxl import load_workbook

    raw = render_xlsx(_ctx(ZtFrameworkCode.CISA_ZTMM_2_0))
    assert raw[:2] == b"PK"
    wb = load_workbook(io.BytesIO(raw))
    assert set(wb.sheetnames) == {"Score Summary", "Answers", "Gap Plan"}


@pytest.mark.unit
def test_dod_xlsx_renders() -> None:
    from openpyxl import load_workbook

    raw = render_xlsx(_ctx(ZtFrameworkCode.DOD_ZTRA))
    wb = load_workbook(io.BytesIO(raw))
    assert "Score Summary" in wb.sheetnames
    ws = wb["Score Summary"]
    # Spot-check the framework cell.
    rows = [(ws.cell(row=r, column=1).value, ws.cell(row=r, column=2).value) for r in range(1, 8)]
    fw = dict(rows).get("Framework")
    assert fw == "DoD ZT Reference Architecture"


@pytest.mark.unit
def test_cisa_xlsx_answers_sheet_row_count() -> None:
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(render_xlsx(_ctx(ZtFrameworkCode.CISA_ZTMM_2_0))))
    ws = wb["Answers"]
    # 37 capabilities + 1 header.
    assert ws.max_row == 38


@pytest.mark.unit
def test_dod_xlsx_answers_sheet_row_count() -> None:
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(render_xlsx(_ctx(ZtFrameworkCode.DOD_ZTRA))))
    ws = wb["Answers"]
    # 50 capabilities + 1 header.
    assert ws.max_row == 51


@pytest.mark.unit
def test_cisa_pdf_renders() -> None:
    raw = render_pdf(_ctx(ZtFrameworkCode.CISA_ZTMM_2_0))
    assert raw.startswith(b"%PDF-")
    assert len(raw) > 2000


def _pdf_text(raw: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(raw))
    return "".join(page.extract_text() for page in reader.pages)


@pytest.mark.unit
def test_cisa_pdf_carries_title_client_framework_and_a_known_pillar() -> None:
    # SMOKE §10: upgrade from %PDF- magic to real content — title, client,
    # framework label, and one known per-pillar rollup row.
    text = _pdf_text(render_pdf(_ctx(ZtFrameworkCode.CISA_ZTMM_2_0)))
    assert "Zero Trust Assessment" in text  # service title
    assert "Atlas Defense Solutions" in text  # client name
    assert "CISA ZTMM 2.0" in text  # framework label
    assert "Identity" in text  # a known CISA pillar rollup row


@pytest.mark.unit
def test_dod_pdf_renders() -> None:
    raw = render_pdf(_ctx(ZtFrameworkCode.DOD_ZTRA))
    assert raw.startswith(b"%PDF-")


@pytest.mark.unit
def test_pdf_handles_empty_gap_list() -> None:
    # Score everyone Optimal (stage 4) and target 3 -> zero gaps.
    ctx = _ctx(ZtFrameworkCode.CISA_ZTMM_2_0, stage=4, target=3)
    raw = render_pdf(ctx)
    assert raw.startswith(b"%PDF-")


@pytest.mark.unit
def test_xlsx_handles_empty_gap_list_with_placeholder() -> None:
    from openpyxl import load_workbook

    ctx = _ctx(ZtFrameworkCode.CISA_ZTMM_2_0, stage=4, target=3)
    wb = load_workbook(io.BytesIO(render_xlsx(ctx)))
    ws = wb["Gap Plan"]
    # Rows shifted by one when #75 added the caption row above the header: the
    # sheet now opens by stating what it is showing and what it is omitting.
    # The placeholder behaviour itself is unchanged, which is the point of this
    # test — only its position moved.
    assert ws.cell(row=1, column=1).value.startswith("All 0 gaps listed.")
    assert ws.max_row == 3
    assert ws.cell(row=3, column=3).value == "No gaps at target stage"
    # The one property the shifted index actually controls, and the reason this
    # test walked past the bug the first time: the italic emphasis must land on
    # the PLACEHOLDER, not on the header the caption pushed into its old slot.
    assert ws.cell(row=3, column=3).font.italic, "placeholder lost its emphasis"
    header = ws.cell(row=2, column=3)
    assert header.value == "Name"
    assert header.font.bold and not header.font.italic, "the header was styled as the placeholder"


@pytest.mark.unit
def test_build_context_falls_back_when_client_none() -> None:
    a, answers = _build_inputs(ZtFrameworkCode.CISA_ZTMM_2_0)
    stage_map = {ans.capability_code: ans.maturity_stage for ans in answers}
    score = compute(ZtFrameworkCode.CISA_ZTMM_2_0, stage_map)
    gap = analyze_gaps(ZtFrameworkCode.CISA_ZTMM_2_0, stage_map)
    ctx = build_context(
        client_legal_name=None,
        service_title="x",
        framework=ZtFrameworkCode.CISA_ZTMM_2_0,
        assessment=a,
        answers=answers,
        score=score,
        gap=gap,
    )
    assert ctx.client_legal_name == "Client"


def _pdf_text(raw: bytes) -> str:
    from pypdf import PdfReader

    return "".join(page.extract_text() for page in PdfReader(io.BytesIO(raw)).pages)


@pytest.mark.unit
def test_every_renderer_discloses_the_gap_plan_truncation() -> None:
    """#75's disclosure in ZT, which nothing asserted.

    Found by the MVP item-9 twin-sweep. `_gap_plan_caption` has two branches and
    only the OTHER one was covered: `test_xlsx_handles_empty_gap_list_with_
    placeholder` asserts `"All 0 gaps listed."`, the branch taken when nothing is
    truncated. The branch that exists *because* of #75 — the one a client sees on
    a real assessment — was asserted by nothing in any renderer, so deleting it
    left the suite green.

    That makes this a worse instance than its ATT&CK twin, which had no caption
    test at all and was therefore visible to a grep. Here a passing test points
    at the wrong branch, so the coverage reads as complete. CSF, the third twin,
    is pinned properly in both its renderers — D-049 pinned one, left one, and
    the camouflage is why nobody noticed.

    Asserted with the literal words and BOTH numbers, per the test-integrity
    gate: `str(n) in blob` would be satisfied by any unrelated occurrence of the
    same digits, and these documents are full of counts.
    """
    ctx = _ctx(ZtFrameworkCode.CISA_ZTMM_2_0, stage=1, target=4)
    shown, total = len(ctx.gap.gaps), ctx.gap.total_gap_count
    assert total > shown, "the fixture does not truncate, so this proves nothing"
    remaining = total - shown
    expected = (
        f"Showing the {shown} highest-priority of {total} gaps; "
        f"{remaining} further gaps not listed."
    )

    from openpyxl import load_workbook

    ws = load_workbook(io.BytesIO(render_xlsx(ctx)))["Gap Plan"]
    assert expected in (ws.cell(row=1, column=1).value or ""), ws.cell(row=1, column=1).value

    from docx import Document

    doc = Document(io.BytesIO(render_docx(ctx)))
    assert expected in "\n".join(p.text for p in doc.paragraphs)

    assert expected in _pdf_text(render_pdf(ctx))


@pytest.mark.unit
def test_the_untruncated_caption_still_says_so_in_every_renderer() -> None:
    """The other branch, in all three rather than one.

    The pre-existing coverage asserted this only in XLSX and only for the
    zero-gap case. A caption that silently stopped rendering in the DOCX or PDF
    would still have passed.
    """
    ctx = _ctx(ZtFrameworkCode.CISA_ZTMM_2_0, stage=4, target=2)
    assert ctx.gap.total_gap_count == len(ctx.gap.gaps), "fixture must NOT truncate"
    expected = f"All {ctx.gap.total_gap_count} gap"

    from openpyxl import load_workbook

    ws = load_workbook(io.BytesIO(render_xlsx(ctx)))["Gap Plan"]
    assert expected in (ws.cell(row=1, column=1).value or "")

    from docx import Document

    doc = Document(io.BytesIO(render_docx(ctx)))
    assert expected in "\n".join(p.text for p in doc.paragraphs)

    assert expected in _pdf_text(render_pdf(ctx))
