"""ATT&CK PDF + XLSX exporter smokes."""

from __future__ import annotations

import io
import uuid

import pytest

from app.attack.analytics import compute as compute_heatmap
from app.attack.catalog import TECHNIQUES
from app.attack.coverage import CoverageStatus
from app.attack.exporters import build_context, render_docx, render_pdf, render_xlsx
from app.models.attack_assessment import (
    AttackAssessment,
    AttackAssessmentStatus,
    AttackCoverage,
)


def _build_inputs(*, default_status: str | None = "covered"):
    a = AttackAssessment(
        id=uuid.uuid4(),
        service_id=uuid.uuid4(),
        version=1,
        status=AttackAssessmentStatus.APPROVED,
    )
    coverage: list[AttackCoverage] = []
    for t in TECHNIQUES:
        coverage.append(
            AttackCoverage(
                id=uuid.uuid4(),
                assessment_id=a.id,
                technique_code=t.id,
                status=default_status,
            )
        )
    coverage_map = {c.technique_code: c.status for c in coverage}
    rollup = compute_heatmap(coverage_map)
    return a, coverage, rollup


def _ctx(*, default_status: str | None = "covered"):
    a, coverage, rollup = _build_inputs(default_status=default_status)
    return build_context(
        client_legal_name="Atlas Defense Solutions",
        service_title="MITRE ATT&CK Coverage",
        assessment=a,
        coverage=coverage,
        rollup=rollup,
    )


@pytest.mark.unit
def test_xlsx_has_three_sheets() -> None:
    from openpyxl import load_workbook

    raw = render_xlsx(_ctx())
    assert raw[:2] == b"PK"
    wb = load_workbook(io.BytesIO(raw))
    assert set(wb.sheetnames) == {"Heatmap Summary", "Coverage", "Gaps"}


@pytest.mark.unit
def test_xlsx_coverage_sheet_has_one_row_per_technique() -> None:
    from openpyxl import load_workbook

    raw = render_xlsx(_ctx())
    wb = load_workbook(io.BytesIO(raw))
    ws = wb["Coverage"]
    assert ws.max_row == len(TECHNIQUES) + 1


@pytest.mark.unit
def test_xlsx_gap_sheet_lists_only_gaps() -> None:
    from openpyxl import load_workbook

    ctx = _ctx(default_status=CoverageStatus.GAP.value)
    raw = render_xlsx(ctx)
    wb = load_workbook(io.BytesIO(raw))
    ws = wb["Gaps"]
    # Header + every technique.
    assert ws.max_row == len(TECHNIQUES) + 1


@pytest.mark.unit
def test_xlsx_gap_placeholder_when_no_gaps() -> None:
    from openpyxl import load_workbook

    ctx = _ctx(default_status=CoverageStatus.COVERED.value)
    raw = render_xlsx(ctx)
    wb = load_workbook(io.BytesIO(raw))
    ws = wb["Gaps"]
    assert ws.max_row == 2  # header + single placeholder
    assert ws.cell(row=2, column=2).value == "No gaps recorded"


@pytest.mark.unit
def test_pdf_renders_valid_bytes() -> None:
    raw = render_pdf(_ctx())
    assert raw.startswith(b"%PDF-")
    assert len(raw) > 2000


def _pdf_text(raw: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(raw))
    return "".join(page.extract_text() for page in reader.pages)


@pytest.mark.unit
def test_pdf_carries_title_client_and_a_known_tactic() -> None:
    # SMOKE §10: upgrade from %PDF- magic to real content — the title prefix
    # (the "&" in "ATT&CK" is reportlab-escaped, so match the stable "MITRE"),
    # the client name, and one known per-tactic rollup row.
    text = _pdf_text(render_pdf(_ctx()))
    assert "MITRE" in text  # service title prefix
    assert "Atlas Defense Solutions" in text  # client name
    assert "Reconnaissance" in text  # a known tactic rollup row


@pytest.mark.unit
def test_pdf_handles_zero_gaps() -> None:
    ctx = _ctx(default_status=CoverageStatus.COVERED.value)
    raw = render_pdf(ctx)
    assert raw.startswith(b"%PDF-")


@pytest.mark.unit
def test_build_context_falls_back_when_client_none() -> None:
    a, coverage, rollup = _build_inputs()
    ctx = build_context(
        client_legal_name=None,
        service_title="x",
        assessment=a,
        coverage=coverage,
        rollup=rollup,
    )
    assert ctx.client_legal_name == "Client"


# ---------------------------------------------------------------------------
# #102: pending review must reach the client-facing document
# ---------------------------------------------------------------------------


def _pending_ctx():
    """Half the catalogue `covered` and withheld, the other half a plain gap.

    This reports **0.0% coverage** -- every positive claim is withheld, so the
    numerator empties while the gaps hold the denominator open. Without the
    pending count beside it, that document is indistinguishable from one saying
    the client owns no controls at all, which is N-033's 607 fabricated gaps
    arriving by a new route.

    So the omission this fixture catches is not a cosmetic one in either
    direction: drop `pending_review` and the same number reads as a finding.
    """
    a, coverage, _ = _build_inputs(default_status=None)
    half = len(coverage) // 2
    for row in coverage[:half]:
        row.status = CoverageStatus.COVERED.value
    for row in coverage[half:]:
        row.status = CoverageStatus.GAP.value
    coverage_map = {c.technique_code: c.status for c in coverage}
    withheld = {c.technique_code for c in coverage[:half]}
    rollup = compute_heatmap(coverage_map, withheld)
    assert rollup.pending_review == half and rollup.covered == 0
    return (
        build_context(
            client_legal_name="Atlas Defense Solutions",
            service_title="MITRE ATT&CK Coverage",
            assessment=a,
            coverage=coverage,
            rollup=rollup,
        ),
        rollup,
    )


@pytest.mark.unit
def test_pdf_states_the_pending_count_beside_the_coverage_percentage() -> None:
    """The percentage is a ratio over what can be CLAIMED, so it is not
    self-describing -- see the analytics test of the same name. The PDF is the
    artifact that reaches the client, so it is the one place a bare percentage
    over withheld rows is a false assurance rather than a UI nit.
    """
    ctx, rollup = _pending_ctx()
    text = _pdf_text(render_pdf(ctx))
    assert f"Pending review {rollup.pending_review}" in text


@pytest.mark.unit
def test_docx_states_the_pending_count_beside_the_coverage_percentage() -> None:
    """The twin. CLAUDE.md: a defect found in one renderer exists in the others
    until checked -- #75 truncated identically in three of them, and #79 exists
    because an earlier change fixed one surface and not its sibling."""
    import io as _io

    from docx import Document

    ctx, rollup = _pending_ctx()
    doc = Document(_io.BytesIO(render_docx(ctx)))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert f"Pending review {rollup.pending_review}" in text


@pytest.mark.unit
def test_every_renderer_carries_pending_review_per_tactic_not_only_overall() -> None:
    """The per-tactic tables, where withholding can push a percentage UP.

    `_pending_ctx` withholds every positive claim, which drives the OVERALL
    figure down to 0% -- an omission there looks obviously wrong. The upward
    direction is the dangerous one and was not covered: a tactic holding one
    confirmed `covered` beside two withheld `partial`s reads 66.7% unwithheld
    and **100%** withheld, so a table without the count states a bare 100% over
    two claims nobody has vouched for.

    XLSX carried the column; the DOCX and PDF twins did not, with nothing saying
    why. That is the unstated-exemption shape CLAUDE.md records for #75/#79, and
    the §14 audit found it here.
    """
    import io as _io

    from docx import Document
    from openpyxl import load_workbook

    ctx, rollup = _pending_ctx()
    assert rollup.pending_review > 0, "the fixture withholds nothing"

    header = "Pending review"
    wb = load_workbook(io.BytesIO(render_xlsx(ctx)))
    ws = wb["Heatmap Summary"]
    tactic_header = next(
        r for r in range(1, ws.max_row + 1) if ws.cell(row=r, column=1).value == "Tactic"
    )
    assert header in [ws.cell(row=tactic_header, column=c).value for c in range(1, 14)]

    doc = Document(_io.BytesIO(render_docx(ctx)))
    docx_tables = [
        [c.text for c in t.rows[0].cells] for t in doc.tables if t.rows and t.rows[0].cells
    ]
    per_tactic = [h for h in docx_tables if h and h[0] == "Tactic"]
    assert per_tactic, "no per-tactic table in the DOCX"
    assert header in per_tactic[0], f"DOCX per-tactic header lacks the count: {per_tactic[0]}"

    # The PDF table is drawn, so assert on extracted text rather than a cell.
    assert header in _pdf_text(render_pdf(ctx))


@pytest.mark.unit
def test_xlsx_states_the_pending_count_and_carries_it_per_tactic() -> None:
    """The third renderer, plus the per-tactic sheet a consultant actually reads."""
    from openpyxl import load_workbook

    ctx, rollup = _pending_ctx()
    wb = load_workbook(io.BytesIO(render_xlsx(ctx)))
    ws = wb["Heatmap Summary"]
    labels = {ws.cell(row=r, column=1).value: ws.cell(row=r, column=2).value for r in range(1, 12)}
    assert labels.get("Pending review") == rollup.pending_review

    header_row = next(
        r for r in range(1, ws.max_row + 1) if ws.cell(row=r, column=1).value == "Tactic"
    )
    headers = [ws.cell(row=header_row, column=col).value for col in range(1, 13)]
    assert "Pending review" in headers


@pytest.mark.unit
def test_the_gap_truncation_disclosure_is_actually_printed() -> None:
    """D-049 rests on "ATT&CK has always disclosed this in its heading" — and
    until now nothing checked.

    Found by the item-3b audit. Both narrative renderers cap the gap list at 50
    and say so; deleting the `(N of M shown)` half of the heading left the whole
    suite green. That is exactly #75's defect — a client-facing truncation with
    no disclosure — sitting unpinned in the service D-049 cited as the good
    example.

    Asserted with the literal words AND both numbers, per the test-integrity
    gate: `str(n) in blob` would be satisfied by any unrelated occurrence of the
    same digits, and this document is full of counts.
    """
    a, coverage, _ = _build_inputs(default_status=CoverageStatus.GAP.value)
    rollup = compute_heatmap({c.technique_code: c.status for c in coverage})
    ctx = build_context(
        client_legal_name="Atlas Defense Solutions",
        service_title="MITRE ATT&CK Coverage",
        assessment=a,
        coverage=coverage,
        rollup=rollup,
    )
    total = rollup.gap
    assert total > 50, "fixture must exceed the cap or this proves nothing"
    expected = f"Top remediation gaps (50 of {total} shown)"

    assert expected in _pdf_text(render_pdf(ctx))

    import io as _io

    from docx import Document

    doc = Document(_io.BytesIO(render_docx(ctx)))
    assert expected in "\n".join(p.text for p in doc.paragraphs)


@pytest.mark.unit
def test_the_xlsx_gap_sheet_does_not_truncate_and_so_makes_no_claim() -> None:
    """The deliberate asymmetry, stated so it is not read as an oversight.

    The workbook is the machine-readable artifact and lists every gap, so it
    carries no "of M shown" heading — there is nothing withheld to disclose.
    """
    from openpyxl import load_workbook

    ctx = _ctx(default_status=CoverageStatus.GAP.value)
    ws = load_workbook(io.BytesIO(render_xlsx(ctx)))["Gaps"]
    assert ws.max_row == len(TECHNIQUES) + 1


@pytest.mark.unit
def test_the_per_technique_sheet_marks_a_withheld_row() -> None:
    """#102's rule reached the summary and stopped at the sheet beside it.

    Found by the item-3b audit. `Heatmap Summary` reported `Covered 0 /
    Pending review N` while the `Coverage` sheet in the SAME workbook listed
    those N techniques as `Covered`, because it printed `cov.status` raw. One
    document, two answers, and the contradiction is visible on adjacent tabs.

    The status still prints — it must, because clearing a citation puts the
    technique back into it. What is added is the column saying the claim is
    being withheld.
    """
    from openpyxl import load_workbook

    ctx, rollup = _pending_ctx()
    wb = load_workbook(io.BytesIO(render_xlsx(ctx)))
    ws = wb["Coverage"]
    headers = [ws.cell(row=1, column=c).value for c in range(1, 9)]
    assert "Pending review" in headers, headers
    col = headers.index("Pending review") + 1
    status_col = headers.index("Status") + 1

    flagged = [
        r
        for r in range(2, ws.max_row + 1)
        if (ws.cell(row=r, column=col).value or "").strip().lower() == "yes"
    ]
    assert (
        len(flagged) == rollup.pending_review
    ), "the sheet disagrees with the summary about how many rows are withheld"
    # And the underlying status survives on those rows, as its rendered label
    # (`coverage_label`), which is what a reader sees.
    assert {ws.cell(row=r, column=status_col).value for r in flagged} == {"Covered"}
